import os
from abc import ABC, abstractmethod
from pathlib import Path

from docx import Document as DocxDocument
from fastmcp import FastMCP
from openpyxl import load_workbook
from pypdf import PdfReader as PyPdfReader
from typing_extensions import override

mcp = FastMCP("Document Reader")


class DocumentReader(ABC):
    """Abstract base class for document readers"""

    @abstractmethod
    def read(self, file_path: str) -> str:
        """Read and extract text from a document"""
        pass


class DocxReader(DocumentReader):
    """DOCX document reader implementation"""

    @override
    def read(self, file_path: str) -> str:
        """Read and extract text from DOCX file"""
        try:
            doc = DocxDocument(file_path)
            text = []

            for paragraph in doc.paragraphs:
                if paragraph.text:
                    text.append(paragraph.text)

            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = " ".join([p.text for p in cell.paragraphs]).strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text.append("\t".join(row_text))

            extracted_text = "\n".join(text)
            return extracted_text if extracted_text else "No text found in the DOCX."
        except Exception as e:
            return f"Error reading DOCX: {str(e)}"


class PdfReader(DocumentReader):
    """PDF document reader implementation"""

    @override
    def read(self, file_path: str) -> str:
        """Read and extract text from PDF file"""
        try:
            with open(file_path, "rb") as file:
                pdf_reader = PyPdfReader(file)
                text = []

                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text.append(page_text.strip())

                extracted_text = "\n\n".join(text)
                return extracted_text if extracted_text else "No text found in the PDF."
        except Exception as e:
            return f"Error reading PDF: {str(e)}"


class TxtReader(DocumentReader):
    """TXT document reader implementation"""

    @override
    def read(self, file_path: str) -> str:
        """Read and extract text from TXT file with encoding handling"""
        encodings = ["utf-8", "gbk", "gb2312", "latin-1"]

        for encoding in encodings:
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    text = f.read()
                return text if text else "No text found in the TXT file."
            except UnicodeDecodeError:
                continue
            except Exception as e:
                return f"Error reading TXT: {str(e)}"

        return "Error reading TXT: Could not decode file with any supported encoding."


class ExcelReader(DocumentReader):
    """Excel document reader implementation"""

    @override
    def read(self, file_path: str) -> str:
        """Read and extract text from Excel file"""
        try:
            wb = load_workbook(file_path, read_only=True)
            text = []

            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                text.append(f"=== Sheet: {sheet_name} ===")

                for row in sheet.iter_rows(values_only=True):
                    row_text = [str(cell) if cell is not None else "" for cell in row]
                    if any(row_text):
                        text.append("\t".join(row_text))

                text.append("")

            extracted_text = "\n".join(text)
            wb.close()
            return (
                extracted_text if extracted_text else "No text found in the Excel file."
            )
        except Exception as e:
            return f"Error reading Excel: {str(e)}"


class DocumentReaderFactory:
    """Factory for creating document readers based on file extension"""

    _readers: dict[str, type[DocumentReader]] = {
        ".txt": TxtReader,
        ".docx": DocxReader,
        ".pdf": PdfReader,
        ".xlsx": ExcelReader,
        ".xls": ExcelReader,
    }

    @classmethod
    def get_reader(cls, file_path: str) -> DocumentReader:
        """Get appropriate reader for the given file"""
        _, ext = os.path.splitext(file_path.lower())
        if ext not in cls._readers:
            raise ValueError(f"Unsupported document type: {ext}")
        return cls._readers[ext]()

    @classmethod
    def is_supported(cls, file_path: str) -> bool:
        """Check if the file type is supported"""
        _, ext = os.path.splitext(file_path.lower())
        return ext in cls._readers


@mcp.tool()
def read_document(filename: str) -> str:
    """
    Reads and extracts text from a specified document file.
    Supports multiple document types: TXT, DOCX, PDF, Excel (XLSX, XLS).

    :param filename: Path to the document file to read
        (supports absolute or relative paths)
    :return: Extracted text from the document
    """
    file_path = Path(filename)

    if not file_path.exists():
        return f"Error: File '{filename}' not found."

    if not DocumentReaderFactory.is_supported(str(file_path)):
        return f"Error: Unsupported document type for file '{filename}'."

    try:
        reader = DocumentReaderFactory.get_reader(str(file_path))
        return reader.read(str(file_path))
    except Exception as e:
        return f"Error reading document: {str(e)}"


def main():
    mcp.run()


if __name__ == "__main__":
    main()
