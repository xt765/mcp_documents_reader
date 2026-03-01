"""生成测试用的示例文档文件。

本脚本创建以下测试文件：
- sample.txt: UTF-8 编码的文本文件
- sample_gbk.txt: GBK 编码的文本文件
- sample.docx: 包含文本和表格的 Word 文档
- sample.pdf: 包含文本的 PDF 文档
- sample.xlsx: 包含多工作表的 Excel 文件
- empty.*: 各种空文档文件
"""

from pathlib import Path

# 安装并导入必要的库
from docx import Document as DocxDocument
from openpyxl import Workbook
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas


def create_fixtures() -> None:
    """创建所有测试文档文件。"""
    fixtures_dir = Path(__file__).parent / "fixtures"
    fixtures_dir.mkdir(parents=True, exist_ok=True)

    # 创建 UTF-8 编码的 TXT 文件
    txt_content = """这是测试文本文件。
包含多行内容。
用于测试 TxtReader。

Special characters: 中文, English, 123, !@#$%
"""
    (fixtures_dir / "sample.txt").write_text(txt_content, encoding="utf-8")
    print("Created: sample.txt")

    # 创建 GBK 编码的 TXT 文件
    gbk_content = "这是 GBK 编码的测试文件。\n包含中文内容。"
    (fixtures_dir / "sample_gbk.txt").write_text(gbk_content, encoding="gbk")
    print("Created: sample_gbk.txt")

    # 创建空 TXT 文件
    (fixtures_dir / "empty.txt").write_text("", encoding="utf-8")
    print("Created: empty.txt")

    # 创建 DOCX 文件
    doc = DocxDocument()
    doc.add_heading("测试文档标题", level=1)
    doc.add_paragraph("这是第一段文本内容。")
    doc.add_paragraph("这是第二段文本内容，包含更多文字。")

    # 添加表格
    table = doc.add_table(rows=3, cols=3)
    table.cell(0, 0).text = "姓名"
    table.cell(0, 1).text = "年龄"
    table.cell(0, 2).text = "城市"
    table.cell(1, 0).text = "张三"
    table.cell(1, 1).text = "25"
    table.cell(1, 2).text = "北京"
    table.cell(2, 0).text = "李四"
    table.cell(2, 1).text = "30"
    table.cell(2, 2).text = "上海"

    doc.save(str(fixtures_dir / "sample.docx"))
    print("Created: sample.docx")

    # 创建空 DOCX 文件
    empty_doc = DocxDocument()
    empty_doc.save(str(fixtures_dir / "empty.docx"))
    print("Created: empty.docx")

    # 创建 PDF 文件
    pdf_path = fixtures_dir / "sample.pdf"
    c = canvas.Canvas(str(pdf_path), pagesize=letter)
    c.setFont("Helvetica", 12)
    c.drawString(100, 750, "This is a test PDF document.")
    c.drawString(100, 730, "It contains multiple lines of text.")
    c.drawString(100, 710, "Used for testing PdfReader.")
    c.showPage()
    c.save()
    print("Created: sample.pdf")

    # 创建空 PDF 文件
    empty_pdf_path = fixtures_dir / "empty.pdf"
    c_empty = canvas.Canvas(str(empty_pdf_path), pagesize=letter)
    c_empty.save()
    print("Created: empty.pdf")

    # 创建 Excel 文件
    wb = Workbook()

    # 第一个工作表
    ws1 = wb.active
    if ws1 is not None:
        ws1.title = "Sheet1"
        ws1["A1"] = "姓名"
        ws1["B1"] = "年龄"
        ws1["C1"] = "城市"
        ws1["A2"] = "张三"
        ws1["B2"] = 25
        ws1["C2"] = "北京"
        ws1["A3"] = "李四"
        ws1["B3"] = 30
        ws1["C3"] = "上海"

    # 第二个工作表
    ws2 = wb.create_sheet("Sheet2")
    ws2["A1"] = "产品"
    ws2["B1"] = "价格"
    ws2["A2"] = "苹果"
    ws2["B2"] = 5.5
    ws2["A3"] = "香蕉"
    ws2["B3"] = 3.2

    wb.save(str(fixtures_dir / "sample.xlsx"))
    print("Created: sample.xlsx")

    # 创建空 Excel 文件
    empty_wb = Workbook()
    empty_active = empty_wb.active
    if empty_active is not None:
        empty_wb.remove(empty_active)  # 删除默认工作表
    empty_wb.create_sheet("Sheet1")  # 创建空工作表
    empty_wb.save(str(fixtures_dir / "empty.xlsx"))
    print("Created: empty.xlsx")

    # 创建损坏的文件用于错误测试
    (fixtures_dir / "corrupted.docx").write_text("This is not a valid docx file")
    print("Created: corrupted.docx")

    (fixtures_dir / "corrupted.pdf").write_text("This is not a valid pdf file")
    print("Created: corrupted.pdf")

    (fixtures_dir / "corrupted.xlsx").write_text("This is not a valid xlsx file")
    print("Created: corrupted.xlsx")

    # 创建二进制文件（无法用任何编码解码）
    binary_data = bytes([0x00, 0x01, 0x02, 0x80, 0x81, 0xFF])
    (fixtures_dir / "binary.txt").write_bytes(binary_data)
    print("Created: binary.txt")

    print("\nAll fixture files created successfully!")


if __name__ == "__main__":
    create_fixtures()
